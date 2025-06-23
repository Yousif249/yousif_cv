# -*- coding: utf-8 -*-
# app.py

import os
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, make_response
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, TextAreaField, SubmitField
from wtforms.validators import DataRequired, Email, Length, EqualTo, ValidationError
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
import json
import logging
import io

# Import for PDF, Word, Excel exports
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

app = Flask(__name__)

# استخدام متغيرات البيئة للتهيئة
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your_super_secret_key_here_for_dev')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///site.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'admin_login'

# تهيئة خط يدعم العربية لـ ReportLab (PDF)
try:
    # استخدام خط Tajawal الموجود في static/css/
    # تأكد من وجود ملف الخط في مجلد static/fonts/
    # يمكنك تحميله من Google Fonts ووضعه في static/fonts/
    pdfmetrics.registerFont(TTFont('Tajawal', 'static/fonts/Tajawal-Regular.ttf'))
    pdfmetrics.registerFont(TTFont('Tajawal-Bold', 'static/fonts/Tajawal-Bold.ttf'))
    ARABIC_FONT = 'Tajawal'
    ARABIC_FONT_BOLD = 'Tajawal-Bold'
except Exception as e:
    app.logger.warning(f"تحذير: لم يتم العثور على خط 'Tajawal-Regular.ttf' أو 'Tajawal-Bold.ttf'. قد لا يتم عرض العربية بشكل صحيح في PDF. يرجى التأكد من توفر الخطوط في مجلد 'static/fonts/'. الخطأ: {e}")
    # fallback to a generic font if custom font fails
    pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
    ARABIC_FONT = 'DejaVuSans'
    app.logger.warning("تحذير: لم يتم العثور على خط 'DejaVuSans.ttf'. قد لا يتم عرض العربية بشكل صحيح في PDF. يرجى التأكد من توفر خط يدعم اللغة العربية (مثل Noto Sans Arabic) في نظامك أو ضمن مجلد التطبيق.")
    pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', 'DejaVuSans-Bold.ttf'))
    ARABIC_FONT_BOLD = 'DejaVuSans-Bold'


# Models
class Admin(db.Model, UserMixin):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(20), unique=True, nullable=False)
    password_hash = db.Column(db.String(128), nullable=False)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class ContactMessage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    phone = db.Column(db.String(20), nullable=False)
    email = db.Column(db.String(100), nullable=False)
    message = db.Column(db.Text, nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

class Visitor(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ip_address = db.Column(db.String(45), nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

# User loader for Flask-Login
@login_manager.user_loader
def load_user(user_id):
    return db.session.get(Admin, int(user_id))

# Forms
class ContactForm(FlaskForm):
    name = StringField('الاسم', validators=[DataRequired('الاسم مطلوب.')])
    phone = StringField('رقم الجوال', validators=[DataRequired('رقم الجوال مطلوب.'), Length(min=8, max=20, message='يجب أن يكون رقم الجوال بين 8 و 20 رقم.')])
    email = StringField('البريد الإلكتروني', validators=[DataRequired('البريد الإلكتروني مطلوب.'), Email('صيغة البريد الإلكتروني غير صحيحة.')])
    message = TextAreaField('الرسالة', validators=[DataRequired('الرسالة مطلوبة.'), Length(min=10, message='يجب أن تكون الرسالة 10 أحرف على الأقل.')])
    submit = SubmitField('إرسال الرسالة')

class AdminLoginForm(FlaskForm):
    username = StringField('اسم المستخدم', validators=[DataRequired('اسم المستخدم مطلوب.')])
    password = PasswordField('كلمة المرور', validators=[DataRequired('كلمة المرور مطلوبة.')])
    submit = SubmitField('تسجيل الدخول')

class EditAdminForm(FlaskForm):
    new_username = StringField('اسم المستخدم الجديد', validators=[DataRequired('اسم المستخدم الجديد مطلوب.')])
    new_password = PasswordField('كلمة المرور الجديدة', validators=[DataRequired('كلمة المرور الجديدة مطلوبة.'), Length(min=6, message='يجب أن تكون كلمة المرور 6 أحرف على الأقل.')])
    confirm_password = PasswordField('تأكيد كلمة المرور الجديدة', validators=[DataRequired('تأكيد كلمة المرور مطلوب.'), EqualTo('new_password', message='كلمتا المرور غير متطابقتين.')])
    submit = SubmitField('تحديث بيانات الأدمن')

# Context processor to make datetime available in all templates
@app.context_processor
def inject_datetime():
    return {'datetime': datetime}

# Before request: create tables and track visitors
@app.before_request
def create_tables_and_track_visitor():
    # Attempt to create tables if they don't exist
    with app.app_context():
        db.create_all()
        # Ensure a default admin exists if no admins are present
        if not Admin.query.first():
            default_admin = Admin(username='admin')
            default_admin.set_password('adminpassword') # يمكنك تغيير كلمة المرور الافتراضية هنا
            db.session.add(default_admin)
            db.session.commit()
            app.logger.info("تم إنشاء حساب الأدمن الافتراضي: المستخدم 'admin'.")

    # Track visitor IP address (simple tracking, no unique constraint to count repeat visits)
    # If you need to count unique visitors per day, you'd add a date component to the query.
    if request.endpoint != 'static': # Avoid tracking requests for static files
        ip_address = request.remote_addr
        # Check if this IP has visited today (basic unique daily visitor check)
        today_start = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
        existing_visitor = Visitor.query.filter_by(ip_address=ip_address).filter(Visitor.timestamp >= today_start).first()
        if not existing_visitor:
            new_visitor = Visitor(ip_address=ip_address)
            db.session.add(new_visitor)
            db.session.commit()
            app.logger.info(f"تم تسجيل زيارة جديدة من IP: {ip_address}")

# Routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/contact', methods=['GET', 'POST'])
def contact():
    form = ContactForm()
    if form.validate_on_submit():
        new_message = ContactMessage(
            name=form.name.data,
            phone=form.phone.data,
            email=form.email.data,
            message=form.message.data
        )
        db.session.add(new_message)
        db.session.commit()
        flash('رسالتك تم إرسالها بنجاح!', 'success')
        return redirect(url_for('contact'))
    return render_template('404.html', form=form) # Note: The template file is named 404.html as per previous conversation, but it functions as the contact page.

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if current_user.is_authenticated:
        return redirect(url_for('admin_dashboard'))
    form = AdminLoginForm()
    if form.validate_on_submit():
        admin = Admin.query.filter_by(username=form.username.data).first()
        if admin and admin.check_password(form.password.data):
            login_user(admin)
            flash('تم تسجيل الدخول بنجاح!', 'success')
            next_page = request.args.get('next')
            return redirect(next_page or url_for('admin_dashboard'))
        else:
            flash('اسم المستخدم أو كلمة المرور غير صحيحة.', 'danger')
    return render_template('admin_login.html', form=form)

@app.route('/admin/dashboard')
@login_required
def admin_dashboard():
    messages = ContactMessage.query.order_by(ContactMessage.timestamp.desc()).all()
    total_visitors = Visitor.query.count()
    return render_template('admin_dashboard.html', messages=messages, total_visitors=total_visitors)

@app.route('/admin/edit', methods=['GET', 'POST'])
@login_required
def edit_admin():
    form = EditAdminForm()
    if form.validate_on_submit():
        current_user.username = form.new_username.data
        current_user.set_password(form.new_password.data)
        db.session.commit()
        flash('تم تحديث بيانات الأدمن بنجاح!', 'success')
        return redirect(url_for('admin_dashboard'))
    # Pre-populate form with current username
    form.new_username.data = current_user.username
    return render_template('edit_admin.html', form=form)

@app.route('/admin/logout')
@login_required
def admin_logout():
    logout_user()
    flash('تم تسجيل الخروج بنجاح.', 'info')
    return redirect(url_for('admin_login'))

# Export Routes
@app.route('/export/excel')
@login_required
def export_excel():
    messages = ContactMessage.query.order_by(ContactMessage.timestamp.desc()).all()
    
    wb = Workbook()
    ws = wb.active
    ws.title = "رسائل الاتصال"

    # Set headers
    headers = ["ID", "الاسم", "الجوال", "البريد الإلكتروني", "الرسالة", "التاريخ"]
    ws.append(headers)

    # Apply bold font to headers
    header_font = Font(bold=True)
    for cell in ws[1]:
        cell.font = header_font
        cell.alignment = Alignment(horizontal='right') # Align headers to right

    # Add data
    for msg in messages:
        ws.append([
            msg.id,
            msg.name,
            msg.phone,
            msg.email,
            msg.message,
            msg.timestamp.strftime('%Y-%m-%d %H:%M:%S')
        ])
    
    # Adjust column widths
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter # Get the column name
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = (max_length + 2) * 1.2
        ws.column_dimensions[column].width = adjusted_width

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    response = make_response(output.getvalue())
    response.headers["Content-Disposition"] = "attachment; filename=contact_messages.xlsx"
    response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return response

@app.route('/export/word')
@login_required
def export_word():
    messages = ContactMessage.query.order_by(ContactMessage.timestamp.desc()).all()
    
    document = Document()
    document.add_heading('رسائل الاتصال', level=1)

    for msg in messages:
        document.add_heading(f'رسالة رقم: {msg.id}', level=2)
        document.add_paragraph(f'الاسم: {msg.name}')
        document.add_paragraph(f'الجوال: {msg.phone}')
        document.add_paragraph(f'البريد الإلكتروني: {msg.email}')
        document.add_paragraph(f'الرسالة: {msg.message}')
        document.add_paragraph(f'التاريخ: {msg.timestamp.strftime("%Y-%m-%d %H:%M:%S")}')
        document.add_paragraph('\n') # Add a newline for separation

    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    response = make_response(output.getvalue())
    response.headers["Content-Disposition"] = "attachment; filename=contact_messages.docx"
    response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return response

@app.route('/export/pdf')
@login_required
def export_pdf():
    messages = ContactMessage.query.order_by(ContactMessage.timestamp.desc()).all()
    
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    p.setFont(ARABIC_FONT, 12)

    # Set initial y position
    y_position = 750
    line_height = 20
    
    p.drawCentredString(letter[0]/2, y_position, "رسائل الاتصال")
    y_position -= (line_height * 2)

    for msg in messages:
        if y_position < 100: # New page if content is too low
            p.showPage()
            p.setFont(ARABIC_FONT, 12)
            y_position = 750
        
        p.setFont(ARABIC_FONT_BOLD, 14)
        p.drawString(50, y_position, f"رسالة رقم: {msg.id}")
        y_position -= line_height

        p.setFont(ARABIC_FONT, 12)
        p.drawString(50, y_position, f"الاسم: {msg.name}")
        y_position -= line_height
        p.drawString(50, y_position, f"الجوال: {msg.phone}")
        y_position -= line_height
        p.drawString(50, y_position, f"البريد الإلكتروني: {msg.email}")
        y_position -= line_height
        
        # Handle long messages
        message_lines = []
        current_line = ""
        words = msg.message.split(' ')
        for word in words:
            if p.stringWidth(current_line + " " + word, ARABIC_FONT, 12) < 500: # Max width for message
                current_line += " " + word
            else:
                message_lines.append(current_line.strip())
                current_line = word
        message_lines.append(current_line.strip())

        p.drawString(50, y_position, "الرسالة:")
        y_position -= line_height
        for line in message_lines:
            if y_position < 100:
                p.showPage()
                p.setFont(ARABIC_FONT, 12)
                y_position = 750
            p.drawString(70, y_position, line) # Indent message lines
            y_position -= line_height
            
        p.drawString(50, y_position, f"التاريخ: {msg.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        y_position -= (line_height * 2) # Extra space after each message

    p.save()

    response = make_response(buffer.getvalue())
    response.headers["Content-Disposition"] = "attachment; filename=contact_messages.pdf"
    response.headers["Content-type"] = "application/pdf"
    return response

if __name__ == '__main__':
    # Initial setup for database and admin user
    with app.app_context():
        db.create_all()
        # Create default admin if not exists
        if not Admin.query.first():
            default_admin = Admin(username='admin')
            default_admin.set_password('adminpassword')
            db.session.add(default_admin)
            db.session.commit()
            app.logger.info("تم إنشاء حساب الأدمن الافتراضي عند بدء التشغيل: المستخدم 'admin'، كلمة المرور 'adminpassword'.")
    
    app.run(debug=True)